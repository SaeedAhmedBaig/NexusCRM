from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "NEXUSCRM_DEVELOPER_GUIDE.md"
DOCX_OUT = ROOT / "NEXUSCRM_DEVELOPER_GUIDE.docx"
PDF_OUT = ROOT / "NEXUSCRM_DEVELOPER_GUIDE.pdf"

BRAND = RGBColor(20, 184, 166)
BRAND_HEX = "#14B8A6"
INK_HEX = "#06110E"


def clean_inline(text):
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = re.sub(r"`([^`]+)`", r"<font name='Courier'>\1</font>", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text


def iter_blocks(markdown):
    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("```"):
            lang = stripped.strip("`").strip()
            code = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            yield ("code", lang, "\n".join(code))
        elif stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            table = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in row):
                    table.append(row)
                i += 1
            i -= 1
            yield ("table", table)
        elif "page-break-before" in stripped or "page-break-after" in stripped:
            yield ("page_break",)
        elif stripped.startswith("<div") or stripped.startswith("</div"):
            pass
        elif stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            yield ("heading", level, stripped[level:].strip())
        elif stripped.startswith("- [ ]"):
            yield ("checkbox", stripped[5:].strip())
        elif stripped.startswith("- "):
            yield ("bullet", stripped[2:].strip())
        elif re.match(r"^\d+\. ", stripped):
            yield ("number", re.sub(r"^\d+\. ", "", stripped))
        elif stripped == "---":
            yield ("rule",)
        elif stripped:
            yield ("paragraph", stripped)
        else:
            yield ("blank",)
        i += 1


def add_docx_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("NexusCRM Developer Guide")
    run.font.size = Pt(9)
    run.font.color.rgb = BRAND


def build_docx(markdown):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    add_docx_footer(section)

    in_cover = True
    for block in iter_blocks(markdown):
        kind = block[0]
        if kind == "page_break":
            doc.add_page_break()
        elif kind == "heading":
            _, level, text = block
            para = doc.add_heading(text, min(level, 4))
            if text.startswith("Chapter ") and level == 1:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if para.runs:
                para.runs[0].font.color.rgb = BRAND if level <= 2 else RGBColor(15, 23, 42)
            in_cover = False
        elif kind == "paragraph":
            para = doc.add_paragraph()
            if in_cover:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", block[1])
            parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith("`") and part.endswith("`"):
                    run = para.add_run(part[1:-1])
                    run.font.name = "Consolas"
                else:
                    para.add_run(part)
        elif kind == "bullet":
            doc.add_paragraph(block[1], style="List Bullet")
        elif kind == "number":
            doc.add_paragraph(block[1], style="List Number")
        elif kind == "checkbox":
            doc.add_paragraph(f"[ ] {block[1]}", style="List Bullet")
        elif kind == "code":
            para = doc.add_paragraph()
            run = para.add_run(block[2])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif kind == "table":
            rows = block[1]
            if rows:
                table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx, cell in enumerate(row):
                        table.cell(r_idx, c_idx).text = re.sub(r"\*\*|`", "", cell)
                        if r_idx == 0:
                            for p in table.cell(r_idx, c_idx).paragraphs:
                                for run in p.runs:
                                    run.bold = True
        elif kind == "rule":
            doc.add_paragraph("")
    doc.save(DOCX_OUT)


def pdf_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor(BRAND_HEX))
    canvas.drawCentredString(LETTER[0] / 2, 0.35 * inch, f"NexusCRM Developer Guide | Page {doc.page}")
    canvas.restoreState()


def build_pdf(markdown):
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="NexusTitle",
            parent=styles["Title"],
            alignment=TA_CENTER,
            textColor=colors.HexColor(BRAND_HEX),
            fontSize=24,
            leading=30,
            spaceAfter=18,
        )
    )
    styles.add(
        ParagraphStyle(
            name="NexusHeading",
            parent=styles["Heading1"],
            textColor=colors.HexColor(BRAND_HEX),
            fontSize=18,
            leading=22,
            spaceBefore=12,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="NexusBody",
            parent=styles["BodyText"],
            fontSize=9,
            leading=12,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="NexusCode",
            parent=styles["Code"],
            fontName="Courier",
            fontSize=7,
            leading=9,
            backColor=colors.HexColor("#F3F4F6"),
            borderPadding=6,
            spaceAfter=8,
        )
    )
    story = []
    for block in iter_blocks(markdown):
        kind = block[0]
        if kind == "page_break":
            story.append(PageBreak())
        elif kind == "heading":
            _, level, text = block
            style = styles["NexusTitle"] if level == 1 else styles["NexusHeading"]
            story.append(Paragraph(clean_inline(text), style))
        elif kind == "paragraph":
            story.append(Paragraph(clean_inline(block[1]), styles["NexusBody"]))
        elif kind == "bullet":
            story.append(Paragraph(f"• {clean_inline(block[1])}", styles["NexusBody"]))
        elif kind == "number":
            story.append(Paragraph(f"• {clean_inline(block[1])}", styles["NexusBody"]))
        elif kind == "checkbox":
            story.append(Paragraph(f"[ ] {clean_inline(block[1])}", styles["NexusBody"]))
        elif kind == "code":
            story.append(Preformatted(block[2], styles["NexusCode"]))
        elif kind == "table":
            rows = block[1]
            if rows:
                max_cols = max(len(r) for r in rows)
                normalized = [r + [""] * (max_cols - len(r)) for r in rows]
                data = [[Paragraph(clean_inline(cell), styles["NexusBody"]) for cell in row] for row in normalized]
                table = Table(data, repeatRows=1)
                table.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(INK_HEX)),
                            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D1D5DB")),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("LEFTPADDING", (0, 0), (-1, -1), 4),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ]
                    )
                )
                story.append(table)
                story.append(Spacer(1, 6))
        elif kind == "rule":
            story.append(Spacer(1, 6))
    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=LETTER,
        rightMargin=0.55 * inch,
        leftMargin=0.55 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.65 * inch,
        title="NexusCRM Developer Guide",
        author="NexusCRM",
    )
    doc.build(story, onFirstPage=pdf_footer, onLaterPages=pdf_footer)


def main():
    markdown = SOURCE.read_text(encoding="utf-8")
    build_docx(markdown)
    build_pdf(markdown)
    print(f"Wrote {DOCX_OUT}")
    print(f"Wrote {PDF_OUT}")


if __name__ == "__main__":
    main()
